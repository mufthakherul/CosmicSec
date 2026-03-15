from datetime import datetime
from pathlib import Path
from typing import List

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel, Field

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    from reportlab.pdfgen import canvas
except Exception:  # pragma: no cover
    canvas = None

import csv
import json

app = FastAPI(title="CosmicSec Report Service", version="1.0.0")


class ReportRequest(BaseModel):
    scan_id: str
    format: str = "pdf"
    findings: List[dict] = Field(default_factory=list)


class ScheduleRequest(BaseModel):
    scan_id: str
    format: str = "pdf"
    cron: str = "0 0 * * *"


REPORT_DIR = Path("/tmp/reports")
REPORT_DIR.mkdir(parents=True, exist_ok=True)
scheduled_jobs: List[dict] = []


def _write_json(file_path: Path, payload: ReportRequest) -> None:
    file_path.write_text(
        json.dumps(
            {
                "scan_id": payload.scan_id,
                "generated_at": datetime.utcnow().isoformat(),
                "findings": payload.findings,
            },
            indent=2,
        ),
        encoding="utf-8",
    )


def _write_csv(file_path: Path, payload: ReportRequest) -> None:
    with file_path.open("w", newline="", encoding="utf-8") as fp:
        writer = csv.DictWriter(fp, fieldnames=["title", "severity", "category", "description"])
        writer.writeheader()
        for row in payload.findings:
            writer.writerow(
                {
                    "title": row.get("title", ""),
                    "severity": row.get("severity", ""),
                    "category": row.get("category", ""),
                    "description": row.get("description", ""),
                }
            )


def _write_html(file_path: Path, payload: ReportRequest) -> None:
    rows = "".join(
        f"<tr><td>{f.get('title','')}</td><td>{f.get('severity','')}</td><td>{f.get('category','')}</td><td>{f.get('description','')}</td></tr>"
        for f in payload.findings
    )
    html = f"""
    <html><body>
    <h1>CosmicSec Report</h1>
    <p>Scan ID: {payload.scan_id}</p>
    <p>Generated: {datetime.utcnow().isoformat()}</p>
    <table border='1' cellspacing='0' cellpadding='5'>
      <tr><th>Title</th><th>Severity</th><th>Category</th><th>Description</th></tr>
      {rows}
    </table>
    </body></html>
    """
    file_path.write_text(html, encoding="utf-8")


def _write_pdf(file_path: Path, payload: ReportRequest) -> None:
    if canvas is None:
        raise HTTPException(status_code=500, detail="reportlab is not installed")

    pdf = canvas.Canvas(str(file_path))
    pdf.setTitle(f"CosmicSec Report {payload.scan_id}")
    y = 800
    pdf.drawString(50, y, f"CosmicSec Report")
    y -= 20
    pdf.drawString(50, y, f"Scan ID: {payload.scan_id}")
    y -= 20
    pdf.drawString(50, y, f"Generated: {datetime.utcnow().isoformat()}")
    y -= 30
    for finding in payload.findings[:25]:
        pdf.drawString(50, y, f"- {finding.get('title', 'untitled')} ({finding.get('severity', 'n/a')})")
        y -= 18
        if y < 80:
            pdf.showPage()
            y = 800
    pdf.save()


def _write_docx(file_path: Path, payload: ReportRequest) -> None:
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx is not installed")

    doc = Document()
    doc.add_heading("CosmicSec Report", level=1)
    doc.add_paragraph(f"Scan ID: {payload.scan_id}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}")
    for finding in payload.findings:
        doc.add_paragraph(f"• {finding.get('title', 'untitled')} ({finding.get('severity', 'n/a')})")
    doc.save(str(file_path))


@app.get("/health")
def health() -> dict:
    return {"status": "healthy", "service": "report-service"}


@app.post("/reports/generate")
def generate_report(payload: ReportRequest) -> dict:
    safe_format = payload.format.lower()
    if safe_format not in {"pdf", "docx", "json", "csv", "html"}:
        safe_format = "pdf"

    out = REPORT_DIR / f"{payload.scan_id}.{safe_format}"
    if safe_format == "json":
        _write_json(out, payload)
    elif safe_format == "csv":
        _write_csv(out, payload)
    elif safe_format == "html":
        _write_html(out, payload)
    elif safe_format == "docx":
        _write_docx(out, payload)
    else:
        _write_pdf(out, payload)

    return {
        "scan_id": payload.scan_id,
        "format": safe_format,
        "artifact": str(out),
        "status": "generated",
        "generated_at": datetime.utcnow().isoformat(),
    }


@app.post("/reports/schedule")
def schedule_report(payload: ScheduleRequest) -> dict:
    job = {
        "scan_id": payload.scan_id,
        "format": payload.format,
        "cron": payload.cron,
        "created_at": datetime.utcnow().isoformat(),
    }
    scheduled_jobs.append(job)
    return {"status": "scheduled", "job": job}
